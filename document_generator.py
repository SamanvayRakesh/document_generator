from openai import OpenAI
from docx import Document
import os

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def generate_document(topic: str, filename: str = "output.docx"):
    prompt = (
        f"Create a professional document on the topic: '{topic}'.\n"
        "Rules:\n"
        "- Include a clear title\n"
        "- Include exactly 5 sections\n"
        "- Each section must have a heading and one concise paragraph\n"
        "- Use neutral, professional language\n\n"
        "Format exactly like this:\n"
        "TITLE:\n"
        "...\n\n"
        "SECTION 1:\n"
        "Heading\n"
        "Content\n\n"
        "SECTION 2:\n"
        "Heading\n"
        "Content\n\n"
        "SECTION 3:\n"
        "Heading\n"
        "Content\n\n"
        "SECTION 4:\n"
        "Heading\n"
        "Content\n\n"
        "SECTION 5:\n"
        "Heading\n"
        "Content\n"
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )

    content = response.choices[0].message.content
    write_docx(content, filename)


def write_docx(text: str, filename: str):
    doc = Document()
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    for line in lines:
        if line.startswith("TITLE:") or line.startswith("SECTION"):
            continue
        if line.isupper():
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)

    doc.save(filename)
